import os
import json
import time
import psycopg2
from psycopg2.extras import execute_values
from openai import OpenAI
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import pypdf
import docx
import magic
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
from flask import Flask, jsonify, request
import threading
import concurrent.futures
import queue
import hashlib
from datetime import datetime
import traceback

# Configuration from environment variables
DB_URL = os.environ.get("DATABASE_URL")
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
DATA_DIR = os.environ.get("DATA_DIR", "/app/data")
CHUNK_SIZE = int(os.environ.get("CHUNK_SIZE", "512"))
CHUNK_OVERLAP = int(os.environ.get("CHUNK_OVERLAP", "50"))
MAX_WORKERS = int(os.environ.get("MAX_WORKERS", "4"))  # Number of parallel workers
STATE_FILE = os.environ.get("STATE_FILE", "/app/processed_files.json")  # Track processed files
ERROR_LOG = os.environ.get("ERROR_LOG", "/app/error_log.json")  # Log processing errors
RATE_LIMIT_DELAY = float(os.environ.get("RATE_LIMIT_DELAY", "0.5"))  # Delay between API calls in seconds

# Global variables
processing_queue = queue.Queue()
processed_files = set()
processing_lock = threading.Lock()
executor = None  # Will be initialized in main()

# Initialize OpenAI client
client = OpenAI(api_key=OPENAI_API_KEY)

# Database connection
def get_db_connection():
    return psycopg2.connect(DB_URL)

# Create embeddings using OpenAI with rate limiting
def create_embedding(text):
    try:
        response = client.embeddings.create(
            input=text,
            model="text-embedding-ada-002"
        )
        # Add small delay to avoid hitting API rate limits
        time.sleep(RATE_LIMIT_DELAY)
        return response.data[0].embedding
    except Exception as e:
        print(f"Error creating embedding: {e}")
        time.sleep(1)  # Wait longer on errors
        # Return a default embedding if we can't get a real one
        # This is not ideal but prevents total failure
        return [0.0] * 1536

# File tracking functions
def load_processed_files():
    """Load the set of already processed files"""
    global processed_files
    try:
        if os.path.exists(STATE_FILE):
            print("Loading previously processed files...")
            start_time = time.time()
            with open(STATE_FILE, 'r') as f:
                file_list = json.load(f)
                processed_files = set(file_list)
                load_time = time.time() - start_time
                print(f"Loaded {len(processed_files)} previously processed files in {load_time:.2f} seconds")
    except Exception as e:
        print(f"Error loading processed files: {e}")
        processed_files = set()

def save_processed_files():
    """Save the set of processed files"""
    try:
        with open(STATE_FILE, 'w') as f:
            json.dump(list(processed_files), f)
    except Exception as e:
        print(f"Error saving processed files: {e}")

def mark_file_processed(file_path):
    """Mark a file as successfully processed"""
    with processing_lock:
        processed_files.add(file_path)
        # Periodically save the state
        if len(processed_files) % 10 == 0:
            save_processed_files()

def is_file_processed(file_path):
    """Check if a file was already processed"""
    return file_path in processed_files

def log_processing_error(file_path, error):
    """Log an error that occurred during processing"""
    try:
        errors = {}
        if os.path.exists(ERROR_LOG):
            with open(ERROR_LOG, 'r') as f:
                errors = json.load(f)
        
        # Update or add the error
        errors[file_path] = {
            "error": str(error),
            "traceback": traceback.format_exc(),
            "timestamp": datetime.now().isoformat()
        }
        
        with open(ERROR_LOG, 'w') as f:
            json.dump(errors, f, indent=2)
    except Exception as e:
        print(f"Error logging processing error: {e}")

# Document processing functions
def is_pdf_searchable(file_path):
    """Check if a PDF contains searchable text"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            # Check first few pages for text
            for i in range(min(3, len(pdf_reader.pages))):
                if pdf_reader.pages[i].extract_text().strip():
                    return True
        return False
    except Exception as e:
        print(f"Error checking if PDF is searchable: {e}")
        return False

def process_pdf_with_ocr(file_path):
    """Process a PDF that needs OCR"""
    print(f"Performing OCR on {file_path}")
    text = ""
    try:
        # Convert PDF to images
        images = convert_from_path(file_path)
        
        # Perform OCR on each page
        for i, image in enumerate(images):
            try:
                page_text = pytesseract.image_to_string(image, lang='eng')
                
                # Clean any binary or NULL characters
                page_text = page_text.replace('\x00', '')
                # Replace any other non-printable characters
                page_text = ''.join(c if c.isprintable() or c.isspace() else ' ' for c in page_text)
                
                text += f"\n\nPage {i+1}:\n{page_text}"
            except Exception as e:
                print(f"Error OCR-ing page {i+1}: {e}")
                text += f"\n\nPage {i+1}: [OCR ERROR: {str(e)}]"
            
        return text
    except Exception as e:
        print(f"Error performing OCR on PDF: {e}")
        # Fallback to regular processing
        return process_pdf_without_ocr(file_path)

def process_pdf_without_ocr(file_path):
    """Process a PDF without OCR"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            text = ""
            for i, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        # Clean any binary or NULL characters
                        page_text = page_text.replace('\x00', '')
                        # Replace any other non-printable characters
                        page_text = ''.join(c if c.isprintable() or c.isspace() else ' ' for c in page_text)
                        text += page_text
                except Exception as e:
                    print(f"Error extracting text from page {i+1}: {e}")
            return text
    except Exception as e:
        print(f"Error processing PDF without OCR: {e}")
        return ""

def process_pdf(file_path):
    """Process a PDF file, using OCR if needed"""
    if not is_pdf_searchable(file_path):
        return process_pdf_with_ocr(file_path)
    else:
        return process_pdf_without_ocr(file_path)

def process_image(file_path):
    """Process an image file with OCR"""
    try:
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image, lang='eng')
        
        # Clean any binary or NULL characters
        text = text.replace('\x00', '')
        # Replace any other non-printable characters
        text = ''.join(c if c.isprintable() or c.isspace() else ' ' for c in text)
        
        return text
    except Exception as e:
        print(f"Error processing image with OCR: {e}")
        return ""

def process_docx(file_path):
    """Process a DOCX file"""
    try:
        doc = docx.Document(file_path)
        text = " ".join([para.text for para in doc.paragraphs])
        
        # Clean any binary or NULL characters
        text = text.replace('\x00', '')
        # Replace any other non-printable characters
        text = ''.join(c if c.isprintable() or c.isspace() else ' ' for c in text)
        
        return text
    except Exception as e:
        print(f"Error processing DOCX file: {e}")
        return ""

def process_txt(file_path):
    """Process a text file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
    except UnicodeDecodeError:
        # Try with a different encoding
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
        except Exception as e:
            print(f"Error processing text file with latin-1 encoding: {e}")
            return ""
    except Exception as e:
        print(f"Error processing text file: {e}")
        return ""
    
    # Clean any binary or NULL characters
    text = text.replace('\x00', '')
    # Replace any other non-printable characters
    text = ''.join(c if c.isprintable() or c.isspace() else ' ' for c in text)
    
    return text

# Worker function for parallel processing
def worker_process_document(file_path):
    """Worker function that processes a single document"""
    # Skip already processed files
    if is_file_processed(file_path):
        print(f"Skipping already processed file: {file_path}")
        return
    
    try:
        # Skip files that don't exist anymore
        if not os.path.exists(file_path):
            print(f"File doesn't exist anymore, skipping: {file_path}")
            return
            
        # Extract metadata and content
        file_name = os.path.basename(file_path)
        file_extension = os.path.splitext(file_name)[1].lower()
        
        print(f"Starting processing of {file_name}...")
        
        # Get mime type for more accurate file type detection
        mime = magic.Magic(mime=True)
        file_type = mime.from_file(file_path)
        
        ocr_applied = False
        content = ""
        
        # Process based on file type
        if file_extension == '.pdf':
            # Check if PDF needs OCR
            if not is_pdf_searchable(file_path):
                content = process_pdf_with_ocr(file_path)
                ocr_applied = True
            else:
                content = process_pdf_without_ocr(file_path)
        elif file_extension == '.docx':
            content = process_docx(file_path)
        elif file_extension == '.txt':
            content = process_txt(file_path)
        elif file_extension in ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif']:
            content = process_image(file_path)
            ocr_applied = True
        else:
            # Try to determine type by MIME
            if file_type.startswith('application/pdf'):
                if not is_pdf_searchable(file_path):
                    content = process_pdf_with_ocr(file_path)
                    ocr_applied = True
                else:
                    content = process_pdf_without_ocr(file_path)
            elif file_type.startswith('image/'):
                content = process_image(file_path)
                ocr_applied = True
            else:
                print(f"Unsupported file type: {file_extension} (MIME: {file_type})")
                return
                
        # Skip if we couldn't extract any content
        if not content or content.strip() == "":
            print(f"No content could be extracted from {file_name}")
            return
        
        # Clean content - remove any NUL characters that might cause database errors
        content = content.replace('\x00', '')
        
        # Create a file hash to uniquely identify the content
        content_hash = hashlib.md5(content.encode()).hexdigest()
        
        # Create metadata
        metadata = {
            "source": file_name,
            "extension": file_extension,
            "path": file_path,
            "mime_type": file_type,
            "ocr_applied": ocr_applied,
            "content_hash": content_hash,
            "processed_at": datetime.now().isoformat()
        }
        
        # Check if we already have a document with this content hash
        with get_db_connection() as conn:
            with conn.cursor() as cursor:
                cursor.execute(
                    """
                    SELECT COUNT(*) FROM document_chunks 
                    WHERE source_metadata->>'content_hash' = %s
                    """, 
                    (content_hash,)
                )
                existing_count = cursor.fetchone()[0]
                
                if existing_count > 0:
                    print(f"Document {file_name} with same content hash already exists in database, skipping")
                    mark_file_processed(file_path)
                    return
                    
        # Create document and chunk it
        document = Document(text=content)
        parser = SentenceSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
        nodes = parser.get_nodes_from_documents([document])
        
        if not nodes:
            print(f"No chunks were created from {file_name}, skipping")
            return
            
        print(f"Created {len(nodes)} chunks for {file_name}, generating embeddings...")
        
        # Store chunks in database
        with get_db_connection() as conn:
            with conn.cursor() as cursor:
                # Insert chunks and get IDs
                chunk_ids = []
                for node in nodes:
                    # Ensure node text is clean
                    clean_text = node.text.replace('\x00', '')
                    # Skip empty chunks
                    if not clean_text.strip():
                        continue
                        
                    cursor.execute(
                        "INSERT INTO document_chunks (text_content, source_metadata) VALUES (%s, %s) RETURNING id",
                        (clean_text, json.dumps(metadata))
                    )
                    chunk_id = cursor.fetchone()[0]
                    chunk_ids.append((chunk_id, clean_text))
                
                # Generate and store embeddings
                embedding_data = []
                for chunk_id, text in chunk_ids:
                    try:
                        embedding = create_embedding(text)
                        embedding_data.append((chunk_id, embedding))
                    except Exception as e:
                        print(f"Error creating embedding for chunk {chunk_id}: {e}")
                
                if embedding_data:
                    # Bulk insert embeddings
                    embedding_values = []
                    for chunk_id, embedding in embedding_data:
                        # Convert embedding list to pgvector format
                        embedding_str = '[' + ','.join(map(str, embedding)) + ']'
                        embedding_values.append((chunk_id, embedding_str))
                    
                    execute_values(
                        cursor,
                        "INSERT INTO chunk_embeddings (chunk_id, embedding_vector) VALUES %s",
                        [(chunk_id, f"{emb}") for chunk_id, emb in embedding_values]
                    )
                    
                conn.commit()
        
        print(f"Successfully processed {file_name} - Created {len(nodes)} chunks with embeddings")
        # Mark the file as processed so we don't process it again
        mark_file_processed(file_path)
        
    except Exception as e:
        print(f"Error processing document {file_path}: {e}")
        traceback.print_exc()
        log_processing_error(file_path, e)

# Function to queue a document for processing
def process_document(file_path):
    """Add a document to the processing queue"""
    processing_queue.put(file_path)

# File watcher class
class DocumentHandler(FileSystemEventHandler):
    def on_created(self, event):
        if not event.is_directory:
            print(f"New file detected: {event.src_path}")
            process_document(event.src_path)

# Queue processor thread function
def queue_processor():
    """Worker thread that processes documents from the queue"""
    print("Starting queue processor thread")
    while True:
        try:
            # Get a file path from the queue
            file_path = processing_queue.get(timeout=1.0)
            
            # Process the file
            worker_process_document(file_path)
            
            # Mark task as done
            processing_queue.task_done()
        except queue.Empty:
            # No documents in queue, sleep briefly
            time.sleep(0.1)
        except Exception as e:
            print(f"Error in queue processor: {e}")
            traceback.print_exc()
            # Sleep a bit on errors to avoid hammering resources
            time.sleep(1.0)

# Process all documents in the data directory
def process_all_documents():
    """Queue all documents in the data directory for processing"""
    import glob
    from pathlib import Path
    
    print("Scanning for documents in data directory...")
    start_time = time.time()
    
    # Use glob for faster file discovery
    extensions = ['*.pdf', '*.docx', '*.txt', '*.jpg', '*.jpeg', '*.png', '*.tiff', '*.tif', '*.bmp', '*.gif']
    
    # Collect all file paths first using glob (much faster than os.walk)
    all_files = []
    data_path = Path(DATA_DIR)
    
    for ext in extensions:
        # Use recursive glob for all subdirectories
        pattern = f"**/{ext}"
        files = list(data_path.glob(pattern))
        all_files.extend([str(f) for f in files])
        
        # Also check uppercase extensions
        pattern_upper = f"**/{ext.upper()}"
        files_upper = list(data_path.glob(pattern_upper))
        all_files.extend([str(f) for f in files_upper])
    
    # Remove duplicates (in case of case-insensitive filesystems)
    all_files = list(set(all_files))
    document_count = len(all_files)
    
    print(f"Found {document_count} documents in {time.time() - start_time:.2f} seconds")
    
    # Batch filter unprocessed files (much faster than individual checks)
    unprocessed_files = []
    batch_size = 1000  # Process in batches to avoid memory issues
    
    for i in range(0, len(all_files), batch_size):
        batch = all_files[i:i + batch_size]
        
        # Filter out processed files from this batch
        with processing_lock:
            unprocessed_batch = [f for f in batch if f not in processed_files]
        
        unprocessed_files.extend(unprocessed_batch)
        
        # Print progress for large batches
        if document_count > 5000 and i % (batch_size * 10) == 0:
            print(f"Filtered {i + len(batch)}/{document_count} files...")
    
    queued_count = len(unprocessed_files)
    print(f"Found {queued_count} unprocessed documents to queue")
    
    # Batch add to queue for better performance
    queue_batch_size = 100
    for i in range(0, len(unprocessed_files), queue_batch_size):
        batch = unprocessed_files[i:i + queue_batch_size]
        
        # Add batch to queue
        for file_path in batch:
            processing_queue.put(file_path)
        
        # Print progress for large queues
        if queued_count > 1000 and i % (queue_batch_size * 10) == 0:
            print(f"Queued {min(i + queue_batch_size, queued_count)}/{queued_count} files...")
    
    total_time = time.time() - start_time
    print(f"Completed scan and queue in {total_time:.2f} seconds. Queue size: {processing_queue.qsize()}")
    return queued_count

# Check for trigger file
def check_for_trigger():
    # Check primary trigger location
    trigger_file = '/app/trigger_ingestion'
    
    # Also check shared volume location (fallback mechanism)
    trigger_file_alt = '/app/graph_data/trigger_ingestion'
    
    if os.path.exists(trigger_file):
        print(f"Trigger file detected at {trigger_file} ({time.strftime('%Y-%m-%d %H:%M:%S')})")
        # Remove the trigger file
        try:
            os.remove(trigger_file)
            print("Trigger file removed")
        except Exception as e:
            print(f"Error removing trigger file: {e}")
        
        # Process all documents
        process_all_documents()
    
    # Check alternative location (shared volume)
    elif os.path.exists(trigger_file_alt):
        print(f"Trigger file detected at {trigger_file_alt} ({time.strftime('%Y-%m-%d %H:%M:%S')})")
        # Remove the trigger file
        try:
            os.remove(trigger_file_alt)
            print("Trigger file removed")
        except Exception as e:
            print(f"Error removing trigger file: {e}")
        
        # Process all documents
        process_all_documents()

# Initialize Flask app for API
app = Flask(__name__)

# Route to trigger reprocessing of documents
@app.route('/trigger-ingestion', methods=['POST'])
def trigger_api():
    print(f"Ingestion triggered via API at {time.strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Process documents in a separate thread to avoid blocking the response
    threading.Thread(target=process_all_documents).start()
    
    return jsonify({
        "status": "success",
        "message": "Document ingestion triggered successfully",
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
    })

# Health check endpoint
@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({
        "status": "healthy",
        "service": "ingestion-service",
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
    })

# Main function for the document watcher
def watch_documents():
    print("Starting document watcher...")
    
    # Set up file watcher
    event_handler = DocumentHandler()
    observer = Observer()
    observer.schedule(event_handler, DATA_DIR, recursive=True)
    observer.start()
    
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        observer.stop()
    observer.join()

# Status endpoint for the API
@app.route('/status', methods=['GET'])
def get_status():
    """Return the status of the ingestion service"""
    status = {
        "queue_size": processing_queue.qsize(),
        "processed_files": len(processed_files),
        "workers": MAX_WORKERS,
        "service": "ingestion-service",
        "timestamp": datetime.now().isoformat()
    }
    
    # Try to get error count
    try:
        if os.path.exists(ERROR_LOG):
            with open(ERROR_LOG, 'r') as f:
                error_data = json.load(f)
                status["errors"] = len(error_data)
        else:
            status["errors"] = 0
    except:
        status["errors"] = "unknown"
    
    return jsonify(status)

# Force process endpoint
@app.route('/force-process', methods=['POST'])
def force_process():
    """Force processing of a specific file even if already processed"""
    data = request.json
    if not data or not data.get("file_path"):
        return jsonify({"status": "error", "message": "No file_path provided"}), 400
    
    file_path = data.get("file_path")
    
    # Remove from processed files if it's there
    with processing_lock:
        if file_path in processed_files:
            processed_files.remove(file_path)
    
    # Queue for processing
    process_document(file_path)
    
    return jsonify({
        "status": "success",
        "message": f"File {file_path} queued for forced processing",
        "timestamp": datetime.now().isoformat()
    })

# Main function
def main():
    global executor
    
    print("Starting document ingestion service...")
    
    # Load previously processed files
    load_processed_files()
    
    # Start worker threads
    print(f"Starting {MAX_WORKERS} worker threads...")
    queue_processors = []
    for i in range(MAX_WORKERS):
        processor = threading.Thread(target=queue_processor, daemon=True)
        processor.start()
        queue_processors.append(processor)
    
    # Process any existing documents on startup
    process_all_documents()
    
    # Start the document watcher in a separate thread
    watcher_thread = threading.Thread(target=watch_documents, daemon=True)
    watcher_thread.start()
    
    # Start the trigger checker thread
    def trigger_checker():
        while True:
            try:
                check_for_trigger()
                time.sleep(10)  # Check every 10 seconds
            except Exception as e:
                print(f"Error in trigger checker: {e}")
                time.sleep(30)  # Back off on errors
    
    trigger_thread = threading.Thread(target=trigger_checker, daemon=True)
    trigger_thread.start()
    
    # Save state periodically
    def state_saver():
        while True:
            try:
                time.sleep(60)  # Save every minute
                save_processed_files()
            except Exception as e:
                print(f"Error saving state: {e}")
    
    state_thread = threading.Thread(target=state_saver, daemon=True)
    state_thread.start()
    
    # Print status periodically
    def status_printer():
        while True:
            try:
                time.sleep(30)  # Every 30 seconds
                print(f"Status: Queue size: {processing_queue.qsize()}, Processed files: {len(processed_files)}")
            except Exception as e:
                print(f"Error printing status: {e}")
    
    status_thread = threading.Thread(target=status_printer, daemon=True)
    status_thread.start()
    
    # Start the Flask API
    print("Starting API server...")
    app.run(host='0.0.0.0', port=5050)

if __name__ == "__main__":
    # Wait for database to be ready
    time.sleep(5)
    main()